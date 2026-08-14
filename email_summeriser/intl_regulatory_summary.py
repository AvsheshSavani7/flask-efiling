"""
International Regulatory Filing Summarizer — Multi-level summaries via Claude API
Usage: python intl_regulatory_summary.py

Handles filings from CADE, EU Commission, CMA, ACCC, SAMR, and other
international competition/antitrust authorities. Extracts structured fields
for jurisdiction, process numbers, review stage, approval status, etc.
"""

from pathlib import Path
from _naming import filing_uid

# ──── PASTE YOUR REGULATORY FILING URL HERE ────
FILING_URL = ""
# ──── OUTPUT FOLDER ────
OUTPUT_DIR = Path(__file__).parent / "Output Summaries"
# ─────────────────────────────────

import os
import sys
import json
import re
import anthropic

try:
    import requests
    from bs4 import BeautifulSoup
    from dotenv import load_dotenv
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install",
                          "requests", "beautifulsoup4", "python-dotenv", "python-docx", "-q"])
    import requests
    from bs4 import BeautifulSoup
    from dotenv import load_dotenv
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load API key from .env (do not sys.exit here — this module is imported by
# the Flask monitor; exiting would kill the gunicorn worker.)
ENV_PATH = Path(__file__).parent.parent / ".env"
load_dotenv(ENV_PATH)


SUMMARY_PROMPT = """You are an expert analyst summarizing international regulatory filings for a merger arbitrage desk.

You receive documents from foreign competition/antitrust authorities (CADE, EU Commission, CMA, ACCC, SAMR, COFECE, KFTC, JFTC, etc.). These may be in any language. Your job is to translate to English, extract structured information, and report ONLY what is NEW.

Given the regulatory filing text below, respond ONLY in valid JSON (no markdown fences) using this schema:

{
  "jurisdiction": "<country or region — e.g., Brazil, EU, UK, Australia, China>",
  "regulatory_body": "<full name of authority — e.g., CADE, European Commission DG Competition, CMA>",
  "document_type": "<e.g., Deficiency Order, Phase 1 Decision, Remedies Proposal, Clearance Decision>",
  "process_number": "<official case/process number as stated in the document>",
  "original_language": "<language the document was written in>",
  "document_date": "<MM/DD/YY>",
  "parties": {
    "acquirer": "<acquiring company name>",
    "target": "<target company name>",
    "tickers": ["<any stock tickers mentioned>"],
    "other_parties": ["<any other named parties — intervenors, complainants, etc.>"]
  },

  "L1_headline": "+ <TICKER or ACQUIRER/TARGET> – <sharpest actionable signal from this update in ~8 words>. | <date>",

  "L2_brief": "<2-3 sentences: what was just filed, what it means, and key context. If review_stage or approval_status changed from prior, state that directly. Do NOT repeat the headline — add context beyond it.>",

  "new_filings": [
    {
      "filing_description": "<exact document name/number that hit the docket>",
      "filing_type_explanation": "<1 sentence: what this document type IS in this jurisdiction's process>",
      "what_it_means": "<1-2 sentences: what this specific filing tells us, stated as fact>",
      "is_first_occurrence": true,
      "source_url": "<URL if available>"
    }
  ],

  "filing_category": "<questionnaire | remedy | objection | intervention | information_request | decision | notification | phase_transition | consultation | extension | administrative | other>",

  "questionnaire_detail": {
    "has_questionnaires": false,
    "respondents": [
      {
        "entity_name": "<company name>",
        "questionnaire_type": "<competitor | customer | supplier | market_participant>",
        "role_description": "<1 sentence: what this entity does in the relevant market>",
        "takeaway": "<1-sentence bottom line: their net stance on the deal's competitive effects. e.g. 'Sees concentrated market; private labels don't discipline prices.' If no response content available, null.>",
        "date_issued": "<date or null>",
        "date_responded": "<date or null>",
        "response_deadline": "<date or null>",
        "topics_covered": ["<topic area 1>", "<topic area 2>"],
        "key_positions_stated": ["<bulleted position/argument if this is a response>"],
        "data_provided": ["<type of data provided, if response>"],
        "concerns_raised": ["<specific competitive concern raised, if any>"],
        "document_references": ["<doc numbers>"],
        "source_url": "<direct link if available>"
      }
    ]
  },

  "remedy_detail": {
    "has_remedies": false,
    "remedy_type": "<structural | behavioral | mixed>",
    "proposed_by": "<parties | authority>",
    "remedies": [
      {
        "description": "<what is being offered/required>",
        "markets_addressed": ["<which competitive concerns this addresses>"],
        "divestiture_assets": "<specific assets/brands/businesses to be divested, if structural>",
        "divestiture_buyer": "<identified buyer or null>",
        "trustee_appointed": "<trustee name or null>",
        "behavioral_commitments": ["<specific behavioral conditions>"],
        "duration": "<time period for behavioral remedies>",
        "deadline": "<date or null>",
        "document_references": [],
        "source_url": "<link if available>"
      }
    ]
  },

  "objection_detail": {
    "has_objections": false,
    "concerns": [
      {
        "theory_of_harm": "<unilateral_effects | coordinated_effects | vertical_foreclosure | conglomerate | other>",
        "markets_affected": ["<product and geographic markets>"],
        "description": "<specific concern as stated by the authority>",
        "evidence_cited": ["<data or analysis referenced>"],
        "parties_implicated": ["<which entities are affected>"],
        "document_references": [],
        "source_url": "<link if available>"
      }
    ]
  },

  "intervention_detail": {
    "has_interventions": false,
    "intervenors": [
      {
        "entity_name": "<company or organization name>",
        "entity_type": "<competitor | customer | supplier | trade_association | government_agency | other>",
        "position": "<supports | opposes | neutral_informational>",
        "key_arguments": ["<argument 1>", "<argument 2>"],
        "remedies_requested": ["<what they want the authority to do, if stated>"],
        "date_filed": "<date>",
        "document_references": [],
        "source_url": "<link if available>"
      }
    ]
  },

  "information_request_detail": {
    "has_info_request": false,
    "requests": [
      {
        "requested_from": "<parties | specific_third_party_name>",
        "request_type": "<deficiency_order | supplemental_request | second_request | other>",
        "data_requested": ["<specific data or documents demanded>"],
        "scope": "<geographic and product market scope of the request>",
        "deadline": "<date>",
        "legal_basis": "<authority/statute for the request>",
        "document_references": [],
        "source_url": "<link if available>"
      }
    ]
  },

  "decision_detail": {
    "has_decision": false,
    "decision_type": "<clearance | conditional_clearance | phase2_escalation | blocking | withdrawal | extension | other>",
    "decision_authority": "<name/title of deciding official>",
    "decision_date": "<date>",
    "legal_basis": "<articles/statutes cited>",
    "conditions_imposed": ["<condition 1>"],
    "effective_date": "<date or null>",
    "appeal_deadline": "<date or null>",
    "vote_or_panel": "<voting breakdown or panel composition if stated>"
  },

  "case_snapshot": {
    "review_stage": "<Phase 1 | Phase 2 | Decision | etc.>",
    "approval_status": "<Approved Unconditionally | Pending - Under Review | etc.>",
    "relevant_markets": [],
    "legal_basis": "...",
    "timeline": { "notification_date": "", "decision_date": "", "response_deadline": "", "next_milestone": "" },
    "related_proceedings": []
  },

  "source_quality": {
    "has_full_document_text": false,
    "sources_unavailable": ["<list URLs that returned SOURCE UNAVAILABLE>"],
    "basis_for_claims": "<'full_text' | 'docket_table_only' | 'email_body_only' | 'partial'>"
  },

  "significance": "<critical | high | medium | low | routine>",
  "significance_reasoning": "<1 sentence>"
}

SIGNIFICANCE RATING GUIDE — rate the update's importance for a merger arbitrage desk:
- "critical": Major milestone that directly affects deal outcome — unconditional approval, approval with conditions, blocking decision, Phase 2 escalation, formal objections issued.
- "high": Substantive progress — formal information request with deadline, remedies proposed or accepted, key competitive concerns identified, significant third-party intervention (e.g., competitor files opposition).
- "medium": Notable development — initial merger notification filed, new jurisdiction opens review, questionnaires issued to market participants, consultation period opens/closes, notable procedural step.
- "low": Minor procedural — internal case routing between departments, routine document acknowledgments, standard protocol registrations, file transfers with no substantive change.
- "routine": No material content — duplicate docket entry, boilerplate filing receipt, purely administrative with zero information value.
When in doubt between two levels, choose the LOWER one — false negatives (missing a low email) are less costly than false positives (sending noise).

Rules:

SOURCE AWARENESS — KNOW WHAT YOU ARE READING:
You will receive text from different source types. You MUST identify which type you have and set source_quality accordingly:

1. DOCKET TABLE (e.g., CADE SEI "Lista de Protocolos"): A list of rows like:
     1755702  |  Despacho Decisório 502  |  20/05/2026  |  CGAA3
     1788319  |  Despacho Decisório 890  |  23/07/2026  |  CGAA3
   This gives you document LABELS, dates, and filing units — NOT document content. You are seeing a table of contents, not the documents themselves. A "Despacho Decisório" (Decision Order) in this table could be anything: a case assignment, a deadline extension, an access grant, a routine procedural step, or yes, an actual clearance decision. You CANNOT tell which from the label alone.
   → Set source_quality.has_full_document_text to false
   → Set source_quality.basis_for_claims to "docket_table_only"

2. FULL DOCUMENT TEXT: Actual paragraphs of reasoning, legal citations in context, explicit outcome language like "aprovação sem restrições" or "approved without conditions" or "application denied". You can read what the authority actually decided and why.
   → Set source_quality.has_full_document_text to true
   → Set source_quality.basis_for_claims to "full_text"

3. EMAIL BODY ONLY: The Hyperion notification email with structured fields (Process, Acquirer, Target, etc.) but no linked document content.
   → Set source_quality.has_full_document_text to false
   → Set source_quality.basis_for_claims to "email_body_only"

4. PARTIAL: Mix — you have some document text but key linked documents returned SOURCE UNAVAILABLE.
   → Set source_quality.has_full_document_text to false
   → Set source_quality.basis_for_claims to "partial"

DECISION EVIDENCE REQUIREMENT — MANDATORY:
You may ONLY set has_decision: true AND decision_type to "clearance", "conditional_clearance", or "blocking" if you have FULL DOCUMENT TEXT (source type 2 above) containing EXPLICIT outcome language. You need actual sentences stating the result — phrases like "aprovação sem restrições", "approved without conditions", "merger cleared", "application denied", "operação aprovada", or equivalent.

When your source is a DOCKET TABLE (source type 1) and you have NO actual document text for a given document:
- Do NOT infer the document's content from its type label
- You MAY note in new_filings that a "Despacho Decisório" was filed, but filing_type_explanation MUST say what it is generically and note that its content was not available
- Set has_decision: false unless you have the actual document text with explicit outcome language

CROSS-REFERENCE REQUIREMENT: For EVERY document you mention in new_filings or L2_brief, you MUST check whether its actual text appears in the source under a "=== LINKED DOCUMENT: <url> ===" section. The source will contain multiple linked document sections — each one is the ACTUAL TEXT of a specific document from the docket.

- If you find the document's text in a LINKED DOCUMENT section: describe what it ACTUALLY SAYS based on what you read. For example, if a Despacho Decisório's text shows it grants access to restricted files ("defiro o pedido de acesso ao processo restrito"), say exactly that — "grants access to restricted case files for the parties' legal representatives."
- If no linked document section contains that document's text: THEN and ONLY THEN say "content not available in the source."
- NEVER claim "content not available" or "specific determinations were not provided" for a document whose actual text you can read in the source. This is a factual error — you have the content in front of you.

NEVER infer approval, denial, or blocking from a document type label. Getting this wrong sends false information to investment professionals.

APPROVAL STATUS — CONSERVATIVE DEFAULT:
The approval_status field in case_snapshot MUST remain "Pending - Under Review" unless you have FULL DOCUMENT TEXT with EXPLICIT clearance/blocking language. A "Despacho Decisório" in a docket table does NOT warrant changing approval_status. When in doubt, keep it as "Pending - Under Review" — a false negative (missing a real approval by one cycle) is infinitely better than a false positive (claiming approval that didn't happen).

WHAT'S NEW — TEMPORAL AWARENESS:
Your ENTIRE response is about what is NEW. If a fact appears in ESTABLISHED FACTS, it is ALREADY KNOWN — do not mention it. The headline must reflect ONLY the new development.

When a docket contains entries spanning multiple dates, identify what TRIGGERED this notification:
- The MOST RECENT docket entries (closest to today/the email date) are the NEW filings — these go in new_filings.
- OLDER entries are HISTORICAL CONTEXT — weave them into L2_brief as narrative context, but do NOT give them their own new_filings entries.
- Example: If an email arrives Aug 10 and the docket has entries from Aug 3 (notification), Aug 4 (orders), Aug 6 (notice), and Aug 10 (DOU publication):
  - new_filings: ONLY the Aug 10 DOU publication (the trigger)
  - L2_brief: "HLAG's USD 4.2B acquisition of ZIM was notified to CADE on 03/08/2026 as an Ordinary Concentration Act. The Superintendência-Geral issued procedural orders on 04/08/2026 granting counsel access. The merger notice (Edital 607) was published in the DOU on 10/08/2026, opening the statutory third-party intervention window."
- The headline and L2_brief should focus on the LATEST development, with earlier filings woven in as narrative background.

WHAT GOES IN new_filings vs L2_brief:
- new_filings: ONLY the filing(s) that TRIGGERED this email — typically the most recent docket entry/entries. This is what the reader needs to act on.
- L2_brief: A narrative summary that provides context. Earlier docket entries (initial notification, procedural orders, access grants, receipts, proxies) are woven in HERE as background, not as separate new_filings cards. If review_stage or approval_status changed from the prior update, state that directly (e.g., "Review stage moved from Phase 1 to Phase 2.").
- For a first-encounter email: the trigger goes in new_filings. Everything before the trigger is narrative context in L2_brief.
- is_first_occurrence on a new_filings entry means this TYPE of filing has never appeared in DEAL HISTORY CONTEXT before.

L2_BRIEF: The L2_brief field is 2-3 sentences that answer: what was just filed, what does it mean, and does it change anything. Do NOT repeat the headline — add context beyond it. For questionnaire responses, highlight the most interesting finding. For decisions, explain the practical consequence. Examples:
  GOOD (if headline is about Santher denial): "Santher had sought interested-party status to monitor remedy negotiations; denial means the merging parties face one fewer intervenor in Phase 1."
  GOOD (if headline is about questionnaire responses): "Among 9 distributor responses, CAMPOS FLORIDOS identified Kenvue as its sole supplier in children's bathing products and flagged 80%+ combined share in feminine absorbents."
  BAD: "CADE issued questionnaires to 40+ market participants." (This just restates the headline.)
  BAD: "Following the initial notification on 05/20/2026, CADE has continued its Phase 1 review by issuing additional questionnaires."

FILING EXPLANATION: filing_type_explanation must be a factual one-sentence definition of what this document type IS. Set it ONLY on the FIRST new_filings entry of each document type. For subsequent entries of the same type, set filing_type_explanation to null — the reader already knows what a Questionário is after the first card.

FILING CATEGORY: Set filing_category to the primary category of this update. Then populate ONLY the matching detail block (set its has_X flag to true). Leave all other detail blocks at has_X: false with empty arrays. Categories: questionnaire, remedy, objection, intervention, information_request, decision, notification, phase_transition, consultation, extension, administrative, other.

NEW FILINGS — ONE PER DOCUMENT (after temporal filtering): For documents that pass the temporal filter above (i.e., the TRIGGERING filings, not historical context), create a SEPARATE new_filings entry for EACH triggering document. Each entry must name the specific entity and document number. But historical/administrative entries from earlier dates belong in L2_brief narrative, NOT in new_filings.

WHAT_IT_MEANS — SUBSTANCE REQUIRED: The what_it_means field on each new_filings entry must state what the filer SAID, ARGUED, or REVEALED — not just that they filed. If the source contains their actual response content, summarize their key position in 1-2 sentences. If the source does NOT contain their response content (restricted file, unavailable), say "Response content not available in source" — do NOT pad with "[Company] filed its questionnaire response on [date]" since that is already stated in filing_description.
  GOOD: "Stated Kenvue is its only supplier among the merging parties; characterized Huggies and Johnson's Baby as non-competing brands for distinct consumer segments."
  GOOD: "Response content not available in source."
  BAD: "PARATY filed its questionnaire response on 07/30/26." (Restates the filing_description — zero information added.)

QUESTIONNAIRE: When filing_category is "questionnaire": set has_questionnaires: true. Create a separate respondent entry for EACH entity named.
- takeaway: ONE sentence that captures this respondent's net stance. The reader should be able to scan 15 takeaways in 30 seconds and know who is concerned and who is not. Examples: "Sees concentrated market; private labels don't discipline prices." / "Has capacity to absorb demand if post-merger prices rise 5-10%." / "No competitive concerns raised." If no response content is available, set to null.
- role_description: WHO is this entity — their business, market position, relationship to the merging parties. This is the reader's first question.
- topics_covered: ONLY set on the FIRST respondent entry or when a respondent's topics DIFFER from the standard questionnaire. Do NOT repeat the same topic list on every respondent — the reader sees it once and understands.
- key_positions_stated: The most important field. What did they ARGUE or REVEAL? Bullet their actual substantive positions — market share claims, supplier relationships, switching costs, competitive dynamics they described. If response content is unavailable, leave empty.
- concerns_raised: Did they flag specific competitive problems? Be concrete.
- For issued questionnaires (not yet responded): capture response_deadline and topics_covered (once).
- Cross-reference ESTABLISHED FACTS to note whether this respondent is NEW or is responding to a previously-issued questionnaire.

REMEDY: When filing_category is "remedy": set has_remedies: true. For each remedy, capture specific assets/brands being divested, behavioral commitments with duration, identified buyers, and which competitive concern each remedy addresses.

OBJECTION: When filing_category is "objection": set has_objections: true. For each competitive concern, state the theory of harm, affected markets, and evidence cited. Use the authority's exact terminology.

INTERVENTION: When filing_category is "intervention": set has_interventions: true. For each intervenor, state their name, type (competitor/customer/etc.), position (supports/opposes), and key arguments as bullets. If they requested specific remedies, list those.

INFORMATION REQUEST: When filing_category is "information_request": set has_info_request: true. Capture what data was demanded, from whom, the scope, deadline, and legal authority for the request.

DECISION: When filing_category is "decision": set has_decision: true. Capture decision type, deciding authority, legal basis, any conditions imposed, effective date, and appeal deadline.

FACTS ONLY: NEVER use: "indicates", "suggests", "signals", "may", "could", "likely", "appears to", "noteworthy that", "significant because". NEVER state what is NOT in the document. If the document is silent on a field, use null or empty array. Every word must trace to a specific phrase in the source document.
  GOOD: "CADE requested revenue data for 2021-2025 across four markets."
  BAD: "The broad scope of information requested indicates potentially detailed competitive analysis ahead."
  GOOD: "The offer expires June 10, 2026."
  BAD: "This tight timeline may create pressure on shareholders to tender quickly."

NO OUTCOME CONFUSION: Never use language that could be misread as claiming an outcome occurred when you are describing a party's REQUEST. A hurried reader scanning the email must not confuse a request with a result.
  BAD: "Parties requested approval without restrictions under summary procedure." (Reads like approval happened.)
  GOOD: "Parties requested that CADE apply the summary (expedited) review procedure." (Clearly a request, not a result.)
  BAD: "requested clearance without conditions" (Reads like clearance happened.)
  GOOD: "sought summary-track review, arguing combined shares fall below CADE's 20% threshold on three of six routes."

FIRST OCCURRENCE: Set is_first_occurrence to true if the DEAL HISTORY CONTEXT has no prior mention of this filing type or event. Set to false if a similar filing/event was previously reported. This tells the reader whether they should have context for this or if it's brand new.

PRECISION: Use the document's exact terminology for legal, regulatory, and financial terms. Do NOT paraphrase in ways that broaden or narrow the stated meaning.

TRANSLATION: Translate all content to English while preserving proper nouns, case numbers, regulatory terms, and legal citations in their original form. For example, keep "Ato de Concentracao" or "Superintendencia-Geral" as-is alongside the English translation.

L1_HEADLINE FORMAT: Must be: + <TICKER or ACQUIRER/TARGET> – <what is NEW in 8 words>. | <date>
- The headline is the ONLY thing a reader sees when scanning 20 emails. It must answer: "Why should I care about THIS update?"
- Capture the THEME — step back and ask: "What is the overall story across ALL the data?" Synthesize the pattern, don't lead with the single most dramatic number.
- For questionnaire responses: What is the balance of opinion? Are most concerned or unconcerned? Where do they agree/disagree? That tension IS the headline.
- Lead with the most important signal, not a procedural description of what was filed.
- BAD: "61 questionnaire responses received from market participants" (procedural, no signal)
- BAD: "CADE issues questionnaires to competitors and customers" (procedural, no signal)
- BAD: "Ever Green estimates 70% post-merger share" (one respondent's view, not the theme)
- BAD: "Competitors flag 70% share, tampon risk" (leads with one stat instead of the overall picture)
- BAD: "Competitors flag ~70% fem-care concentration; baby-bath draws fewer objections" (still anchored on one respondent's number — use the PATTERN not the stat)
- GOOD: "Competitors split — fem-care concentration concerns vs. sufficient private-label constraint" (captures the tension)
- GOOD: "Broad third-party pushback on fem-care overlap; baby-bath draws fewer objections" (pattern across markets)
- GOOD: "Phase 2 opened; authority cites unresolved market-share concerns" (consequence, not procedure)
- GOOD: "Unconditional approval — no remedies required" (outcome)
- If the filing contains substantive findings/positions/concerns, headline the overall pattern. If the filing is purely procedural (notification accepted, questionnaires issued), then describe the procedural step but keep it tight.
- 8 words is a target, not a hard limit — clarity beats brevity.

Extract exact case/process numbers, legal article citations, official names. Identify the review stage and approval status precisely. Note any deadlines or response requirements with exact dates.

REGULATORY FILING TEXT:
"""

EXTRACTION_GUIDANCE = """This is an international regulatory/competition authority filing (e.g., CADE, EU Commission, CMA, ACCC, SAMR).
Extract VERBATIM in the original language (preserve exact case numbers, legal citations, party names, regulatory terms):
- DEAL HISTORY CONTEXT section (if present) — pass through verbatim
- Authority name and jurisdiction
- Case/process number and document type
- Document type names in original language with English translation
- Filing/docket registration numbers and dates
- Parties involved (acquirer, target, other named parties)
- Action taken by the authority (decision, information request, approval, etc.)
- Review stage and approval status
- Markets identified (product markets, geographic markets)
- Information or documents requested (for deficiency orders)
- Conditions or remedies imposed or proposed
- Competitive concerns or theories of harm stated
- Legal basis (statutes, articles, regulations cited)
- All dates (notification, decision, deadlines, next milestones)
- Questionnaire respondent names, entity types, dates issued/responded, topics, deadlines
- References to parallel proceedings in other jurisdictions
- Any facts relevant to deal timeline or closing
- Names and titles of signing officials
- IMPORTANT: Identify the source type. CADE SEI pages have two formats:
  (a) Process page ("Lista de Protocolos"): A TABLE of rows with columns [Document Number | Document Type | Date | Unit]. Example: "1788319 | Despacho Decisório 890 | 23/07/2026 | CGAA3". This is a docket INDEX — you can see what was filed but NOT what it says. Flag this: "NOTE: Source is a CADE SEI docket listing — contains document labels and dates only, not document content."
  (b) Document page ("Consulta Externa"): The actual text of a specific document — legal reasoning, decisions with explicit outcome language, orders with specific instructions. This IS document content.
- Do NOT infer outcomes (approved, denied, blocked) from document type labels in a docket table. "Despacho Decisório" means "Decision Order" but could be any procedural order — case assignments, deadline extensions, access grants, or internal routing. Only actual document text tells you what was decided."""


def fetch_filing_text(source: str) -> str:
    """Fetch and extract text from an international regulatory filing.

    Tries the standard fetch pipeline first. If the source URL returns
    403 or 503 (common for bot-protected government sites), falls back
    to Jina Reader (r.jina.ai) for extraction.
    """
    from fetch_utils import fetch_text, extract_relevant_sections

    is_url = isinstance(source, str) and source.startswith("http")

    if is_url:
        try:
            full_text = fetch_text(source, word_limit=0)
        except requests.exceptions.HTTPError as e:
            if e.response is not None and e.response.status_code in (403, 503):
                print(f"   Got {e.response.status_code} — trying Jina Reader fallback...")
                jina_url = f"https://r.jina.ai/{source}"
                headers = {"User-Agent": "MergerArbDashboard/1.0 (merger-arb-research@outlook.com)"}
                resp = requests.get(jina_url, headers=headers, timeout=60)
                resp.raise_for_status()
                full_text = resp.text
            else:
                raise
    else:
        full_text = fetch_text(source, word_limit=0)

    return extract_relevant_sections(full_text, EXTRACTION_GUIDANCE)


def summarize(text: str, model: str = "claude-opus-4-8") -> dict:
    """Call Claude API to produce structured summary."""
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY_TEST"))

    msg = client.messages.create(
        model=model,
        max_tokens=16384,
        messages=[{
            "role": "user",
            "content": SUMMARY_PROMPT + "\n\n" + text
        }]
    )

    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    # If response was truncated, try to close the JSON
    if msg.stop_reason != "end_turn":
        if raw.count('{') > raw.count('}'):
            raw += '"' + '}' * (raw.count('{') - raw.count('}'))
        if raw.count('[') > raw.count(']'):
            raw += ']' * (raw.count('[') - raw.count(']'))

    return json.loads(raw)


def print_summary(s: dict):
    """Pretty-print the structured summary."""
    print("\n" + "=" * 70)
    print(f"  INTERNATIONAL REGULATORY FILING SUMMARY")
    print("=" * 70)

    print(f"\n   Jurisdiction:    {s.get('jurisdiction', 'N/A')}")
    print(f"   Authority:       {s.get('regulatory_body', 'N/A')}")
    print(f"   Document Type:   {s.get('document_type', 'N/A')}")
    print(f"   Process Number:  {s.get('process_number', 'N/A')}")
    print(f"   Language:        {s.get('original_language', 'N/A')}")
    print(f"   Date:            {s.get('document_date', 'N/A')}")

    parties = s.get("parties", {})
    print(f"\n   Acquirer:  {parties.get('acquirer', 'N/A')}")
    print(f"   Target:    {parties.get('target', 'N/A')}")
    if parties.get("tickers"):
        print(f"   Tickers:   {', '.join(parties['tickers'])}")
    if parties.get("other_parties"):
        print(f"   Other:     {', '.join(parties['other_parties'])}")

    print(f"\n   L1: HEADLINE")
    print(f"   {s.get('L1_headline', s.get('headline', 'N/A'))}")

    if s.get("L2_brief") or s.get("what_happened"):
        print(f"\n   L2: BRIEF")
        print(f"   {s.get('L2_brief', s.get('what_happened', ''))}")

    # New filings (trigger filings only)
    new_filings = s.get("new_filings", [])
    if new_filings:
        print(f"\n   FILINGS")
        for f in new_filings:
            first = " [FIRST REPORTED]" if f.get("is_first_occurrence") else ""
            print(f"     * {f.get('filing_description', 'N/A')}{first}")
            if f.get("filing_type_explanation"):
                print(f"       Type: {f['filing_type_explanation']}")
            if f.get("what_it_means"):
                print(f"       Meaning: {f['what_it_means']}")
            if f.get("source_url"):
                print(f"       Link: {f['source_url']}")

    print(f"\n   Filing Category: {s.get('filing_category', 'N/A')}")

    # Questionnaire detail
    qd = s.get("questionnaire_detail", {})
    if qd.get("has_questionnaires"):
        print(f"\n   QUESTIONNAIRE DETAIL")
        for r in qd.get("respondents", []):
            print(f"     Entity: {r.get('entity_name', 'N/A')} ({r.get('questionnaire_type', 'N/A')})")
            if r.get("role_description"):
                print(f"       Role: {r['role_description']}")
            if r.get("date_issued"):
                print(f"       Issued: {r['date_issued']}")
            if r.get("date_responded"):
                print(f"       Responded: {r['date_responded']}")
            if r.get("response_deadline"):
                print(f"       Deadline: {r['response_deadline']}")
            for t in r.get("topics_covered", []):
                print(f"       Topic: {t}")
            for p in r.get("key_positions_stated", []):
                print(f"       Position: {p}")
            for c in r.get("concerns_raised", []):
                print(f"       Concern: {c}")
            if r.get("source_url"):
                print(f"       Link: {r['source_url']}")

    # Remedy detail
    rd = s.get("remedy_detail", {})
    if rd.get("has_remedies"):
        print(f"\n   REMEDY DETAIL ({rd.get('remedy_type', 'N/A')}, proposed by {rd.get('proposed_by', 'N/A')})")
        for rem in rd.get("remedies", []):
            print(f"     * {rem.get('description', 'N/A')}")
            if rem.get("divestiture_assets"):
                print(f"       Assets: {rem['divestiture_assets']}")
            if rem.get("divestiture_buyer"):
                print(f"       Buyer: {rem['divestiture_buyer']}")
            for bc in rem.get("behavioral_commitments", []):
                print(f"       Commitment: {bc}")
            if rem.get("duration"):
                print(f"       Duration: {rem['duration']}")

    # Objection detail
    od = s.get("objection_detail", {})
    if od.get("has_objections"):
        print(f"\n   OBJECTION DETAIL")
        for c in od.get("concerns", []):
            print(f"     Theory: {c.get('theory_of_harm', 'N/A')}")
            print(f"     Description: {c.get('description', 'N/A')}")
            for m in c.get("markets_affected", []):
                print(f"       Market: {m}")

    # Intervention detail
    ivd = s.get("intervention_detail", {})
    if ivd.get("has_interventions"):
        print(f"\n   INTERVENTION DETAIL")
        for iv in ivd.get("intervenors", []):
            print(f"     {iv.get('entity_name', 'N/A')} ({iv.get('entity_type', 'N/A')}) — {iv.get('position', 'N/A')}")
            for arg in iv.get("key_arguments", []):
                print(f"       - {arg}")

    # Information request detail
    ird = s.get("information_request_detail", {})
    if ird.get("has_info_request"):
        print(f"\n   INFORMATION REQUEST DETAIL")
        for req in ird.get("requests", []):
            print(f"     From: {req.get('requested_from', 'N/A')} ({req.get('request_type', 'N/A')})")
            for d_item in req.get("data_requested", []):
                print(f"       - {d_item}")
            if req.get("deadline"):
                print(f"       Deadline: {req['deadline']}")

    # Decision detail
    dd = s.get("decision_detail", {})
    if dd.get("has_decision"):
        print(f"\n   DECISION DETAIL")
        print(f"     Type: {dd.get('decision_type', 'N/A')}")
        if dd.get("decision_authority"):
            print(f"     Authority: {dd['decision_authority']}")
        if dd.get("decision_date"):
            print(f"     Date: {dd['decision_date']}")
        if dd.get("legal_basis"):
            print(f"     Legal Basis: {dd['legal_basis']}")
        for cond in dd.get("conditions_imposed", []):
            print(f"     Condition: {cond}")

    # Case snapshot
    cs = s.get("case_snapshot", {})
    print(f"\n   CASE SNAPSHOT")
    print(f"   Review Stage:     {cs.get('review_stage', 'N/A')}")
    print(f"   Approval Status:  {cs.get('approval_status', 'N/A')}")
    if cs.get("relevant_markets"):
        print(f"   Relevant Markets:")
        for m in cs["relevant_markets"]:
            print(f"     - {m}")
    if cs.get("legal_basis"):
        print(f"   Legal Basis:      {cs['legal_basis']}")
    timeline = cs.get("timeline", {})
    if any(v for v in timeline.values()):
        print(f"   Timeline:")
        for k, v in timeline.items():
            if v:
                label = k.replace("_", " ").title()
                print(f"     - {label}: {v}")
    if cs.get("related_proceedings"):
        print(f"   Related Proceedings:")
        for r in cs["related_proceedings"]:
            print(f"     - {r}")

    # Significance
    print(f"\n   Significance: {s.get('significance', 'N/A')}")
    if s.get("significance_reasoning"):
        print(f"   Reason: {s['significance_reasoning']}")

    print("=" * 70)


def export_docx(s: dict, filepath: str = None):
    """Export summary to a formatted Word document."""
    from fetch_utils import is_empty_value, has_content, add_field

    jurisdiction = s.get("jurisdiction", "INTL")
    authority = s.get("regulatory_body", "")
    date = s.get("document_date", "")
    parties = s.get("parties", {})
    acquirer = parties.get("acquirer", "UNKNOWN")

    if filepath is None:
        safe_acquirer = re.sub(r'[^\w\-\.]', '_', acquirer)
        safe_jurisdiction = re.sub(r'[^\w\-\.]', '_', jurisdiction)
        filepath = f"IntlReg_{safe_jurisdiction}_{safe_acquirer}_{date.replace('/', '-')}.docx"

    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Regulatory Filing: {jurisdiction} — {authority}", level=0)
    title.runs[0].font.size = Pt(18)

    # Metadata
    meta = doc.add_paragraph()
    add_field(meta, "Document Type: ", s.get("document_type"), newline=False)
    meta.add_run("    ")
    add_field(meta, "Process Number: ", s.get("process_number"), newline=False)

    meta2 = doc.add_paragraph()
    add_field(meta2, "Date: ", date, newline=False)
    meta2.add_run("    ")
    add_field(meta2, "Original Language: ", s.get("original_language"), newline=False)

    meta3 = doc.add_paragraph()
    add_field(meta3, "Acquirer: ", parties.get("acquirer"), newline=False)
    meta3.add_run("    ")
    add_field(meta3, "Target: ", parties.get("target"), newline=False)
    if parties.get("tickers"):
        meta3.add_run("    ")
        add_field(meta3, "Tickers: ", ", ".join(parties["tickers"]), newline=False)

    cs = s.get("case_snapshot", {})

    # L1: Headline
    doc.add_heading("Headline", level=1)
    p = doc.add_paragraph()
    headline_text = s.get("L1_headline", s.get("headline", ""))
    run = p.add_run(headline_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 68)

    # L2: Brief
    brief = s.get("L2_brief", s.get("what_happened", ""))
    if brief:
        doc.add_heading("Brief", level=1)
        doc.add_paragraph(brief)

    # Filings
    new_filings = s.get("new_filings", [])
    if new_filings:
        doc.add_heading("Filings", level=1)
        for f in new_filings:
            p = doc.add_paragraph()
            run = p.add_run(f.get("filing_description", ""))
            run.bold = True
            if f.get("is_first_occurrence"):
                run2 = p.add_run("  [FIRST REPORTED]")
                run2.font.color.rgb = RGBColor(0, 102, 68)
                run2.font.size = Pt(9)
            if not is_empty_value(f.get("filing_type_explanation")):
                p2 = doc.add_paragraph()
                run = p2.add_run(f["filing_type_explanation"])
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            if not is_empty_value(f.get("what_it_means")):
                doc.add_paragraph(f["what_it_means"])

    # Filing category
    if not is_empty_value(s.get("filing_category")):
        p = doc.add_paragraph()
        add_field(p, "Filing Category: ", s.get("filing_category"), newline=False)

    # Questionnaire detail
    qd = s.get("questionnaire_detail", {})
    if qd.get("has_questionnaires"):
        doc.add_heading("Questionnaire Detail", level=1)
        for r in qd.get("respondents", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{r.get('entity_name', 'N/A')} ({r.get('questionnaire_type', '')})")
            run.bold = True
            if not is_empty_value(r.get("role_description")):
                doc.add_paragraph(r["role_description"])
            info_p = doc.add_paragraph()
            if r.get("date_issued"):
                add_field(info_p, "Issued: ", r["date_issued"], newline=True)
            if r.get("date_responded"):
                add_field(info_p, "Responded: ", r["date_responded"], newline=True)
            if r.get("response_deadline"):
                add_field(info_p, "Deadline: ", r["response_deadline"], newline=True)
            for t in r.get("topics_covered", []):
                if not is_empty_value(t):
                    doc.add_paragraph(f"Topic: {t}", style="List Bullet")
            for pos in r.get("key_positions_stated", []):
                if not is_empty_value(pos):
                    doc.add_paragraph(f"Position: {pos}", style="List Bullet")
            for c in r.get("concerns_raised", []):
                if not is_empty_value(c):
                    doc.add_paragraph(f"Concern: {c}", style="List Bullet")

    # Remedy detail
    rd = s.get("remedy_detail", {})
    if rd.get("has_remedies"):
        doc.add_heading("Remedy Detail", level=1)
        p = doc.add_paragraph()
        add_field(p, "Type: ", rd.get("remedy_type"), newline=False)
        p.add_run("    ")
        add_field(p, "Proposed by: ", rd.get("proposed_by"), newline=False)
        for rem in rd.get("remedies", []):
            p = doc.add_paragraph()
            run = p.add_run(rem.get("description", ""))
            run.bold = True
            if not is_empty_value(rem.get("divestiture_assets")):
                doc.add_paragraph(f"Assets: {rem['divestiture_assets']}", style="List Bullet")
            if not is_empty_value(rem.get("divestiture_buyer")):
                doc.add_paragraph(f"Buyer: {rem['divestiture_buyer']}", style="List Bullet")
            for bc in rem.get("behavioral_commitments", []):
                if not is_empty_value(bc):
                    doc.add_paragraph(f"Commitment: {bc}", style="List Bullet")

    # Objection detail
    od = s.get("objection_detail", {})
    if od.get("has_objections"):
        doc.add_heading("Objection Detail", level=1)
        for c in od.get("concerns", []):
            p = doc.add_paragraph()
            run = p.add_run(f"Theory: {c.get('theory_of_harm', 'N/A')}")
            run.bold = True
            if not is_empty_value(c.get("description")):
                doc.add_paragraph(c["description"])
            for m in c.get("markets_affected", []):
                if not is_empty_value(m):
                    doc.add_paragraph(f"Market: {m}", style="List Bullet")

    # Intervention detail
    ivd = s.get("intervention_detail", {})
    if ivd.get("has_interventions"):
        doc.add_heading("Intervention Detail", level=1)
        for iv in ivd.get("intervenors", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{iv.get('entity_name', 'N/A')} ({iv.get('entity_type', '')}) — {iv.get('position', '')}")
            run.bold = True
            for arg in iv.get("key_arguments", []):
                if not is_empty_value(arg):
                    doc.add_paragraph(arg, style="List Bullet")

    # Information request detail
    ird = s.get("information_request_detail", {})
    if ird.get("has_info_request"):
        doc.add_heading("Information Request Detail", level=1)
        for req in ird.get("requests", []):
            p = doc.add_paragraph()
            run = p.add_run(f"From: {req.get('requested_from', 'N/A')} ({req.get('request_type', '')})")
            run.bold = True
            for d_item in req.get("data_requested", []):
                if not is_empty_value(d_item):
                    doc.add_paragraph(d_item, style="List Bullet")
            if not is_empty_value(req.get("deadline")):
                doc.add_paragraph(f"Deadline: {req['deadline']}")

    # Decision detail
    dd = s.get("decision_detail", {})
    if dd.get("has_decision"):
        doc.add_heading("Decision Detail", level=1)
        p = doc.add_paragraph()
        add_field(p, "Decision Type: ", dd.get("decision_type"), newline=True)
        add_field(p, "Authority: ", dd.get("decision_authority"), newline=True)
        add_field(p, "Date: ", dd.get("decision_date"), newline=True)
        add_field(p, "Legal Basis: ", dd.get("legal_basis"), newline=True)
        for cond in dd.get("conditions_imposed", []):
            if not is_empty_value(cond):
                doc.add_paragraph(f"Condition: {cond}", style="List Bullet")

    # Case Snapshot
    doc.add_heading("Case Snapshot", level=1)
    status_para = doc.add_paragraph()
    add_field(status_para, "Review Stage: ", cs.get("review_stage"), newline=True)
    add_field(status_para, "Approval Status: ", cs.get("approval_status"), newline=True)
    add_field(status_para, "Legal Basis: ", cs.get("legal_basis"), newline=True)

    items = cs.get("relevant_markets")
    if has_content(items):
        doc.add_heading("Relevant Markets", level=2)
        for m in items:
            if not is_empty_value(m):
                doc.add_paragraph(m, style="List Bullet")

    timeline = cs.get("timeline", {})
    if has_content(timeline):
        doc.add_heading("Timeline", level=2)
        for k, v in timeline.items():
            if not is_empty_value(v):
                label = k.replace("_", " ").title()
                p = doc.add_paragraph()
                add_field(p, f"{label}: ", v, newline=False)

    items = cs.get("related_proceedings")
    if has_content(items):
        doc.add_heading("Related Proceedings", level=2)
        for r in items:
            if not is_empty_value(r):
                doc.add_paragraph(r, style="List Bullet")

    # Significance
    p = doc.add_paragraph()
    add_field(p, "Significance: ", s.get("significance"), newline=True)
    add_field(p, "Reason: ", s.get("significance_reasoning"), newline=False)

    doc.save(filepath)
    return filepath


def main():
    source = FILING_URL

    print(f"Fetching international regulatory filing from: {source}")

    text = fetch_filing_text(source)
    print(f"Extracted {len(text.split())} words of text")

    print("Generating summary via Claude Opus...")
    result = summarize(text)

    print_summary(result)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Save JSON
    jurisdiction = result.get("jurisdiction", "INTL")
    safe_jurisdiction = re.sub(r'[^\w\-\.]', '_', jurisdiction)
    uid = filing_uid(FILING_URL) if "sec.gov" in str(FILING_URL) else re.sub(r'[^\w]', '', str(FILING_URL))[-8:]
    out_path = OUTPUT_DIR / f"intl_reg_summary_{safe_jurisdiction}_{uid}.json"
    out_path.write_text(json.dumps(result, indent=2))
    print(f"\nRaw JSON saved to: {out_path}")

    # Save DOCX
    parties = result.get("parties", {})
    acquirer = parties.get("acquirer", "UNKNOWN")
    safe_acquirer = re.sub(r'[^\w\-\.]', '_', acquirer)
    date = result.get("document_date", "")
    safe_date = date.replace("/", "-")
    filename = f"IntlReg_{safe_jurisdiction}_{safe_acquirer}_{safe_date}_{uid}.docx"
    docx_path = export_docx(result, str(OUTPUT_DIR / filename))
    print(f"Word doc saved to:  {docx_path}")

    return result


if __name__ == "__main__":
    if not os.getenv("ANTHROPIC_API_KEY_TEST"):
        print(f"ANTHROPIC_API_KEY not found. Check your .env at:\n   {ENV_PATH}")
        sys.exit(1)
    main()
