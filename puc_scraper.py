import asyncio
import requests
import time
import zipfile
import os
import tempfile
import urllib3
import subprocess
import platform
from io import BytesIO
from urllib.parse import urljoin, urlparse
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError
from playwright_stealth import Stealth
from bs4 import BeautifulSoup
import PyPDF2

# Try to import openpyxl for Excel file parsing
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("Warning: openpyxl not installed. XLSX files will not be parsed. Install with: pip install openpyxl")

# Try to import python-docx for Word document parsing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX files will not be parsed. Install with: pip install python-docx")

# Disable SSL warnings for requests with verify=False
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def extract_text_from_old_doc(file_path):
    """
    Extract text from old binary .doc files (pre-2007 format).
    Uses platform-specific tools. Prioritizes antiword for Linux servers.

    Args:
        file_path: Path to the .doc file

    Returns:
        Extracted text as string, or None if extraction fails
    """
    try:
        system = platform.system()

        # Try antiword first (works on Linux, macOS, Windows)
        # This is the primary method for Render/Linux servers
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
                stderr=subprocess.DEVNULL
            )
            if result.returncode == 0 and result.stdout.strip():
                print(f"Successfully extracted text using antiword")
                return result.stdout.strip()
        except FileNotFoundError:
            print("antiword not found, trying platform-specific methods...")
        except Exception as e:
            print(f"antiword failed: {str(e)}")

        # Platform-specific fallbacks
        if system == "Darwin":  # macOS
            # Use textutil (built-in on macOS)
            try:
                result = subprocess.run(
                    ["textutil", "-convert", "txt", "-stdout", file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    print(f"Successfully extracted text using textutil (macOS)")
                    return result.stdout.strip()
            except Exception as e:
                print(f"textutil failed: {str(e)}")

        elif system == "Linux":
            # Try LibreOffice as fallback for Linux
            try:
                # Create temp directory for output
                with tempfile.TemporaryDirectory() as temp_dir:
                    subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "txt:Text",
                         "--outdir", temp_dir, file_path],
                        capture_output=True,
                        timeout=30,
                        stderr=subprocess.DEVNULL
                    )
                    # Read the converted file
                    txt_file = os.path.join(temp_dir, os.path.splitext(
                        os.path.basename(file_path))[0] + ".txt")
                    if os.path.exists(txt_file):
                        with open(txt_file, 'r', encoding='utf-8') as f:
                            extracted = f.read().strip()
                            if extracted:
                                print(
                                    f"Successfully extracted text using LibreOffice")
                                return extracted
            except FileNotFoundError:
                print("LibreOffice not found")
            except Exception as e:
                print(f"LibreOffice failed: {str(e)}")

        print(f"All extraction methods failed for {file_path}")
        return None

    except Exception as e:
        print(f"Error extracting text from old .doc file: {str(e)}")
        return None


def is_old_doc_format(file_content):
    """
    Check if a file is in old binary .doc format (pre-2007).

    Args:
        file_content: Binary content of the file

    Returns:
        True if it's an old .doc format, False otherwise
    """
    # Old .doc files start with D0 CF 11 E0 (OLE2 compound document)
    return file_content[:8] == b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'


TARGET_URL = "https://ipinfo.io/ip"
API_KEY = "1cc2815b0dfbc0b37d0218bc5f4325d1"

# proxy_host = "95.135.111.121"
# proxy_port = 49815
# proxy_username = "OcFu9LdAZ1XbXpa"
# proxy_password = "T7jFlurhMqMWhha"


proxy_host = "108.59.242.138"
proxy_port = 46885
proxy_username = "GSenAgrfKhuNWkd"
proxy_password = "8lmVa5yl0pKp9MI"


CAPTCHA_SOLVER_URL = "http://2captcha.com/in.php"
CAPTCHA_RESULT_URL = "http://2captcha.com/res.php"


async def get_sitekey(page):
    content = await page.content()
    soup = BeautifulSoup(content, "html.parser")
    sitekey = None
    iframe = soup.find(
        "iframe", src=lambda x: x and "challenges.cloudflare.com/cdn-cgi/challenge-platform" in x)
    if iframe and "sitekey=" in iframe["src"]:
        import urllib.parse
        q = urllib.parse.parse_qs(iframe["src"].split("?", 1)[1])
        sitekey = q.get("sitekey", [None])[0]
    if not sitekey:
        inp = soup.find("input", {"name": "cf-turnstile-sitekey"})
        if inp:
            sitekey = inp["value"]
    if not sitekey:
        div = soup.find("div", {"data-sitekey": True})
        if div:
            sitekey = div.get("data-sitekey")
    return sitekey


async def playwright_2captcha_fetch_puc(url, wait_time=15):
    async with Stealth().use_async(async_playwright()) as p:
        browser = await p.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--ignore-certificate-errors",
                f"--proxy-server=http://{proxy_host}:{proxy_port}"
            ]
        )
        context = await browser.new_context(
            proxy={
                "server": f"http://{proxy_host}:{proxy_port}",
                "username": proxy_username,
                "password": proxy_password
            },
            ignore_https_errors=True,
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            extra_http_headers={
                "Accept-Language": "en-US,en;q=0.9",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8"
            }

        )
        page = await context.new_page()
        try:
            await page.goto(url, timeout=120_000)
        except PlaywrightTimeoutError:
            print("Timed out waiting for initial page load. Saving what loaded.")
            html = await page.content()
            with open("debug_puc_initial_timeout.html", "w", encoding="utf-8") as f:
                f.write(html)
            await browser.close()
            return html
        await asyncio.sleep(5)

        sitekey = await get_sitekey(page)
        if not sitekey:
            print("No Turnstile found. Returning current page.")
            await asyncio.sleep(20)
            html = await page.content()
            await browser.close()
            return html

        print(f"Turnstile sitekey found: {sitekey}")

        # 2Captcha submission
        data = {
            "key": API_KEY,
            "method": "turnstile",
            "sitekey": sitekey,
            "pageurl": url,
            "json": 1,
        }
        resp = requests.post(CAPTCHA_SOLVER_URL, data=data)
        task_id = resp.json().get("request")
        if not task_id:
            print("2Captcha task creation failed:", resp.text)
            await browser.close()
            return ""

        print(f"2Captcha task id: {task_id} - waiting for solution...")
        token = None
        for _ in range(30):
            time.sleep(5)
            params = {
                "key": API_KEY,
                "action": "get",
                "id": task_id,
                "json": 1,
            }
            r = requests.get(CAPTCHA_RESULT_URL, params=params)
            result = r.json()
            if result["status"] == 1:
                token = result["request"]
                break
            elif result["request"] != "CAPCHA_NOT_READY":
                print("2Captcha error:", result)
                break
            print("Waiting for captcha solution ...")
        if not token:
            print("Captcha was not solved in time.")
            await browser.close()
            return ""

        print(f"Captcha solved: {token[:8]}... (truncated)")

        # DEBUG: Save HTML just before submitting the form
        html_before = await page.content()
        with open("debug_puc_before_submit.html", "w", encoding="utf-8") as f:
            f.write(html_before)

        # 3. Inject token into correct form and submit
        await page.evaluate("""
            (token) => {
                let input = document.querySelector('input[name="cf-turnstile-response"]');
                if (!input) {
                    input = document.createElement('input');
                    input.type = "hidden";
                    input.name = "cf-turnstile-response";
                    document.body.appendChild(input);
                }
                input.value = token;
                let form = input.form || document.querySelector('form');
                if (form) {
                    form.submit();
                } else {
                    location.reload();
                }
            }
        """, token)

        # 4. Wait for next page
        try:
            await page.wait_for_load_state('networkidle', timeout=50000)
        except Exception:
            print("Waiting extra for challenge navigation...")
            await asyncio.sleep(10)

        for _ in range(10):
            if not page.is_closed():
                try:
                    html_after = await page.content()
                    break
                except Exception:
                    await asyncio.sleep(2)
            else:
                return ""
        with open("debug_puc_after_submit.html", "w", encoding="utf-8") as f:
            f.write(html_after)

        print("Saving HTML ...")
        await browser.close()
        return html_after


def extract_metadata_from_html(html_content):
    """
    Extract metadata from HTML paragraphs with strong tags.

    Args:
        html_content: HTML content as string

    Returns:
        Dictionary containing metadata key-value pairs
    """
    soup = BeautifulSoup(html_content, "html.parser")
    metadata = {}

    # Find all paragraphs with strong tags
    paragraphs = soup.find_all("p")

    for p in paragraphs:
        strong = p.find("strong")
        if strong:
            # Extract the key from strong tag
            key = strong.get_text(strip=True)

            # Extract the value (everything after the strong tag)
            # Remove the strong tag text and any nbsp/whitespace
            full_text = p.get_text(separator=" ", strip=True)
            # Remove the key from the beginning
            value = full_text.replace(key, "", 1).strip()

            # Clean up common HTML entities
            value = value.replace("\xa0", " ").strip()

            if key and value:
                metadata[key] = value

    return metadata


def extract_zip_links_from_html(html_content, base_url):
    """
    Extract ZIP file links from the HTML table. If no ZIP files found,
    extract PDF and DOC files instead.

    Args:
        html_content: HTML content as string
        base_url: Base URL to resolve relative links

    Returns:
        Dictionary containing zip_files list, doc_files list, and metadata dict
    """
    soup = BeautifulSoup(html_content, "html.parser")
    zip_files = []
    doc_files = []

    # Extract metadata from paragraphs
    metadata = extract_metadata_from_html(html_content)

    # Find the table with documents
    table = soup.find("table", class_="table")
    if not table:
        print("No table found in HTML")
        return {"zip_files": zip_files, "doc_files": doc_files, "metadata": metadata}

    # Find all rows in tbody (skip header)
    tbody = table.find("tbody")
    if tbody:
        rows = tbody.find_all("tr")
    else:
        # If no tbody, get all rows except the first (header)
        rows = table.find_all("tr")[1:]  # Skip header row

    for row in rows:
        cells = row.find_all("td")
        if len(cells) < 3:
            continue

        # Get the link in the first cell
        link_elem = cells[0].find("a")
        if not link_elem or not link_elem.get("href"):
            continue

        file_url = link_elem.get("href")
        name = link_elem.get_text(strip=True)
        description = cells[1].get_text(strip=True) if len(cells) > 1 else ""
        file_type = cells[2].get_text(strip=True) if len(cells) > 2 else ""

        # Resolve relative URLs
        if not file_url.startswith("http"):
            file_url = urljoin(base_url, file_url)

        # Process ZIP files
        if file_type.upper() == "ZIP" or file_url.upper().endswith(".ZIP"):
            zip_files.append({
                "url": file_url,
                "name": name,
                "description": description,
                "type": file_type
            })
        # Process PDF and DOC files
        elif (file_type.upper() in ["PDF", "DOC", "DOCX"] or
              file_url.upper().endswith((".PDF", ".DOC", ".DOCX"))):
            doc_files.append({
                "url": file_url,
                "name": name,
                "description": description,
                "type": file_type if file_type else file_url.split('.')[-1].upper()
            })

    # If no ZIP files found, log that we're using doc files
    if not zip_files and doc_files:
        print(
            f"No ZIP files found. Found {len(doc_files)} PDF/DOC file(s) instead.")

    return {"zip_files": zip_files, "doc_files": doc_files, "metadata": metadata}


def download_and_extract_zip(zip_url, output_dir=None):
    """
    Download a ZIP file and extract its contents.

    Args:
        zip_url: URL of the ZIP file to download
        output_dir: Directory to extract files to (if None, uses temp directory)

    Returns:
        Dictionary with extraction results
    """
    try:
        # Use proxy if available
        proxy_dict = None
        if proxy_host and proxy_port:
            proxy_dict = {
                "http": f"http://{proxy_username}:{proxy_password}@{proxy_host}:{proxy_port}",
                "https": f"http://{proxy_username}:{proxy_password}@{proxy_host}:{proxy_port}"
            }

        # Download the ZIP file
        # Disable SSL verification to match Playwright browser context settings
        print(f"Downloading ZIP from: {zip_url}")
        response = requests.get(
            zip_url, proxies=proxy_dict, timeout=120, verify=False)
        response.raise_for_status()

        # Create output directory if not provided
        if output_dir is None:
            output_dir = tempfile.mkdtemp(prefix="puc_extracted_")
        else:
            os.makedirs(output_dir, exist_ok=True)

        # Save ZIP file temporarily
        zip_filename = os.path.basename(urlparse(zip_url).path)
        zip_path = os.path.join(output_dir, zip_filename)

        with open(zip_path, "wb") as f:
            f.write(response.content)

        print(f"ZIP file saved to: {zip_path}")

        # Extract ZIP file
        extracted_files = []
        extract_dir = os.path.join(output_dir, f"{zip_filename}_extracted")
        os.makedirs(extract_dir, exist_ok=True)

        with zipfile.ZipFile(zip_path, "r") as zip_ref:
            zip_ref.extractall(extract_dir)
            extracted_files = zip_ref.namelist()

        print(f"Extracted {len(extracted_files)} files to: {extract_dir}")

        # Get full paths of extracted files and extract text from PDFs
        extracted_file_paths = []
        pdf_text_content = []

        for file_name in extracted_files:
            full_path = os.path.join(extract_dir, file_name)
            if os.path.isfile(full_path):
                file_info = {
                    "name": file_name,
                    "path": full_path,
                    "size": os.path.getsize(full_path),
                    "type": "unknown"
                }

                # Check if it's a PDF file and extract text
                if file_name.lower().endswith('.pdf'):
                    file_info["type"] = "pdf"
                    try:
                        # Read PDF file and extract text
                        with open(full_path, 'rb') as pdf_file:
                            pdf_content = pdf_file.read()

                            # Validate PDF header
                            if pdf_content.startswith(b'%PDF'):
                                pdf_reader = PyPDF2.PdfReader(
                                    BytesIO(pdf_content))
                                text = ""
                                page_count = len(pdf_reader.pages)

                                for page in pdf_reader.pages:
                                    page_text = page.extract_text()
                                    text += page_text + "\n"

                                pdf_text = text.strip()
                                file_info["pdf_text"] = pdf_text
                                file_info["pdf_page_count"] = page_count
                                file_info["pdf_text_length"] = len(pdf_text)

                                pdf_text_content.append({
                                    "file_name": file_name,
                                    "page_count": page_count,
                                    "text_length": len(pdf_text),
                                    "text": pdf_text
                                })

                                print(
                                    f"Extracted text from PDF {file_name}: {len(pdf_text)} characters, {page_count} pages")
                            else:
                                print(
                                    f"Invalid PDF file {file_name}: doesn't start with PDF header")
                                file_info["pdf_error"] = "Invalid PDF header"
                    except Exception as e:
                        print(
                            f"Error extracting text from PDF {file_name}: {str(e)}")
                        file_info["pdf_error"] = str(e)

                # Check if it's an XLSX file and extract text
                elif file_name.lower().endswith(('.xlsx', '.xls')):
                    file_info["type"] = "xlsx"
                    try:
                        if not OPENPYXL_AVAILABLE:
                            file_info["xlsx_error"] = "openpyxl library not installed"
                            print(
                                f"Skipping XLSX {file_name}: openpyxl not available")
                        else:
                            # Read and parse Excel file
                            workbook = openpyxl.load_workbook(
                                full_path, data_only=True)
                            text_content = []

                            # Process all sheets
                            for sheet_name in workbook.sheetnames:
                                sheet = workbook[sheet_name]
                                sheet_text = [f"Sheet: {sheet_name}"]
                                sheet_text.append("=" * 50)

                                # Read all rows
                                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                                    # Filter out completely empty rows
                                    if any(cell is not None and str(cell).strip() for cell in row):
                                        # Convert row to text, handling None values
                                        row_text = []
                                        for cell in row:
                                            if cell is None:
                                                row_text.append("")
                                            else:
                                                # Convert cell to string and handle special characters
                                                cell_str = str(cell).strip()
                                                row_text.append(cell_str)

                                        # Join with tab or pipe separator for readability
                                        sheet_text.append(" | ".join(row_text))

                                text_content.append("\n".join(sheet_text))
                                # Add spacing between sheets
                                text_content.append("\n")

                            xlsx_text = "\n".join(text_content).strip()
                            file_info["xlsx_text"] = xlsx_text
                            file_info["xlsx_sheet_count"] = len(
                                workbook.sheetnames)
                            file_info["xlsx_sheets"] = workbook.sheetnames
                            file_info["xlsx_text_length"] = len(xlsx_text)

                            print(
                                f"Extracted text from XLSX {file_name}: {len(xlsx_text)} characters, {len(workbook.sheetnames)} sheet(s)")
                            workbook.close()
                    except Exception as e:
                        print(
                            f"Error extracting text from XLSX {file_name}: {str(e)}")
                        file_info["xlsx_error"] = str(e)

                # Check if it's a DOCX file and extract text
                elif file_name.lower().endswith(('.docx', '.doc')):
                    file_info["type"] = "docx" if file_name.lower().endswith(
                        '.docx') else "doc"
                    try:
                        # Check if it's an old .doc format
                        with open(full_path, 'rb') as f:
                            file_header = f.read(8)

                        is_old_doc = is_old_doc_format(file_header)

                        if is_old_doc:
                            # Handle old binary .doc format
                            print(
                                f"Detected old binary .doc format for {file_name}")
                            docx_text = extract_text_from_old_doc(full_path)

                            if docx_text:
                                file_info["docx_text"] = docx_text
                                file_info["docx_text_length"] = len(docx_text)
                                file_info["doc_format"] = "binary (pre-2007)"
                                print(
                                    f"Extracted text from old DOC {file_name}: {len(docx_text)} characters")
                            else:
                                file_info["docx_error"] = "Failed to extract text from old .doc format"
                                print(
                                    f"Failed to extract text from old DOC {file_name}")
                        else:
                            # Handle new .docx format
                            if not DOCX_AVAILABLE:
                                file_info["docx_error"] = "python-docx library not installed"
                                print(
                                    f"Skipping DOCX {file_name}: python-docx not available")
                            else:
                                # Read and parse Word document
                                doc = docx.Document(full_path)
                                text_content = []

                                # Extract text from all paragraphs
                                for paragraph in doc.paragraphs:
                                    if paragraph.text.strip():
                                        text_content.append(
                                            paragraph.text.strip())

                                # Also extract text from tables if present
                                for table in doc.tables:
                                    text_content.append("\n[Table]")
                                    for row in table.rows:
                                        row_text = []
                                        for cell in row.cells:
                                            cell_text = cell.text.strip()
                                            row_text.append(cell_text)
                                        if any(row_text):
                                            text_content.append(
                                                " | ".join(row_text))
                                    text_content.append("[End Table]\n")

                                docx_text = "\n".join(text_content).strip()
                                file_info["docx_text"] = docx_text
                                file_info["docx_text_length"] = len(docx_text)
                                file_info["docx_paragraph_count"] = len(
                                    [p for p in doc.paragraphs if p.text.strip()])
                                file_info["docx_table_count"] = len(doc.tables)
                                file_info["doc_format"] = "Office Open XML (.docx)"

                                print(
                                    f"Extracted text from DOCX {file_name}: {len(docx_text)} characters, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
                    except Exception as e:
                        print(
                            f"Error extracting text from DOC/DOCX {file_name}: {str(e)}")
                        file_info["docx_error"] = str(e)

                extracted_file_paths.append(file_info)

        # Return simplified file list with only essential info
        simplified_files = []
        for file_info in extracted_file_paths:
            simplified_file = {
                "name": file_info["name"],
                "type": file_info.get("type", "unknown")
            }

            # Add text content if it's a PDF
            if file_info.get("type") == "pdf" and file_info.get("pdf_text"):
                simplified_file["text"] = file_info["pdf_text"]

            # Add text content if it's an XLSX file
            elif file_info.get("type") == "xlsx" and file_info.get("xlsx_text"):
                simplified_file["text"] = file_info["xlsx_text"]

            # Add text content if it's a DOCX file
            elif file_info.get("type") == "docx" and file_info.get("docx_text"):
                simplified_file["text"] = file_info["docx_text"]

            simplified_files.append(simplified_file)

        if pdf_text_content:
            print(f"Extracted text from {len(pdf_text_content)} PDF file(s)")

        return {
            "success": True,
            "files": simplified_files
        }

    except Exception as e:
        print(f"Error downloading/extracting ZIP from {zip_url}: {str(e)}")
        return {
            "success": False,
            "zip_url": zip_url,
            "error": str(e)
        }


def download_and_extract_document(doc_url, file_type):
    """
    Download an individual document (PDF or DOC) and extract its text content.

    Args:
        doc_url: URL of the document to download
        file_type: Type of file (PDF, DOC, DOCX)

    Returns:
        Dictionary with extraction results containing file info and extracted text
    """
    try:
        # Use proxy if available
        proxy_dict = None
        if proxy_host and proxy_port:
            proxy_dict = {
                "http": f"http://{proxy_username}:{proxy_password}@{proxy_host}:{proxy_port}",
                "https": f"http://{proxy_username}:{proxy_password}@{proxy_host}:{proxy_port}"
            }

        # Download the file
        print(f"Downloading {file_type} from: {doc_url}")
        response = requests.get(
            doc_url, proxies=proxy_dict, timeout=120, verify=False)
        response.raise_for_status()

        file_content = response.content
        file_name = os.path.basename(urlparse(doc_url).path)

        result = {
            "name": file_name,
            "type": file_type.lower(),
            "url": doc_url,
            "size": len(file_content)
        }

        # Extract text based on file type
        if file_type.upper() == "PDF" or file_name.lower().endswith('.pdf'):
            try:
                # Validate PDF header
                if file_content.startswith(b'%PDF'):
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                    text = ""
                    page_count = len(pdf_reader.pages)

                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        text += page_text + "\n"

                    pdf_text = text.strip()
                    result["text"] = pdf_text
                    result["page_count"] = page_count
                    result["text_length"] = len(pdf_text)

                    print(
                        f"Extracted text from PDF {file_name}: {len(pdf_text)} characters, {page_count} pages")
                else:
                    result["error"] = "Invalid PDF header"
                    print(
                        f"Invalid PDF file {file_name}: doesn't start with PDF header")
            except Exception as e:
                result["error"] = str(e)
                print(f"Error extracting text from PDF {file_name}: {str(e)}")

        elif file_type.upper() in ["DOC", "DOCX"] or file_name.lower().endswith(('.doc', '.docx')):
            try:
                # Check if it's an old .doc format
                is_old_doc = is_old_doc_format(file_content)

                if is_old_doc:
                    # Handle old binary .doc format
                    print(f"Detected old binary .doc format for {file_name}")

                    # Save to temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp_file:
                        tmp_file.write(file_content)
                        tmp_path = tmp_file.name

                    try:
                        docx_text = extract_text_from_old_doc(tmp_path)

                        if docx_text:
                            result["text"] = docx_text
                            result["text_length"] = len(docx_text)
                            result["doc_format"] = "binary (pre-2007)"
                            print(
                                f"Extracted text from old DOC {file_name}: {len(docx_text)} characters")
                        else:
                            result["error"] = "Failed to extract text from old .doc format"
                            print(
                                f"Failed to extract text from old DOC {file_name}")
                    finally:
                        # Clean up temp file
                        os.unlink(tmp_path)
                else:
                    # Handle new .docx format
                    if not DOCX_AVAILABLE:
                        result["error"] = "python-docx library not installed"
                        print(
                            f"Skipping DOCX {file_name}: python-docx not available")
                    else:
                        # Save to temp file for docx library
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                            tmp_file.write(file_content)
                            tmp_path = tmp_file.name

                        try:
                            # Read and parse Word document
                            doc = docx.Document(tmp_path)
                            text_content = []

                            # Extract text from all paragraphs
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    text_content.append(paragraph.text.strip())

                            # Also extract text from tables if present
                            for table in doc.tables:
                                text_content.append("\n[Table]")
                                for row in table.rows:
                                    row_text = []
                                    for cell in row.cells:
                                        cell_text = cell.text.strip()
                                        row_text.append(cell_text)
                                    if any(row_text):
                                        text_content.append(
                                            " | ".join(row_text))
                                text_content.append("[End Table]\n")

                            docx_text = "\n".join(text_content).strip()
                            result["text"] = docx_text
                            result["text_length"] = len(docx_text)
                            result["paragraph_count"] = len(
                                [p for p in doc.paragraphs if p.text.strip()])
                            result["table_count"] = len(doc.tables)
                            result["doc_format"] = "Office Open XML (.docx)"

                            print(
                                f"Extracted text from DOCX {file_name}: {len(docx_text)} characters")
                        finally:
                            # Clean up temp file
                            os.unlink(tmp_path)
            except Exception as e:
                result["error"] = str(e)
                print(
                    f"Error extracting text from DOC/DOCX {file_name}: {str(e)}")

        return {
            "success": True,
            "file": result
        }

    except Exception as e:
        print(
            f"Error downloading/extracting document from {doc_url}: {str(e)}")
        return {
            "success": False,
            "url": doc_url,
            "error": str(e)
        }


def process_puc_documents(html_content, base_url, extract_zips=True):
    """
    Process PUC documents from HTML: extract ZIP links, download and extract them.
    If no ZIP files found, process PDF and DOC files instead.

    Args:
        html_content: HTML content as string
        base_url: Base URL of the page
        extract_zips: Whether to download and extract files

    Returns:
        Dictionary with zip_urls array, doc_urls array, extracted_files array, and metadata dict
    """
    # Extract ZIP links, document links, and metadata from HTML
    extraction_result = extract_zip_links_from_html(html_content, base_url)
    zip_files = extraction_result["zip_files"]
    doc_files = extraction_result["doc_files"]
    metadata = extraction_result["metadata"]

    zip_urls = [zip_info["url"] for zip_info in zip_files]
    doc_urls = [doc_info["url"] for doc_info in doc_files]

    print(f"Found {len(zip_files)} ZIP file(s) in the table")
    print(f"Found {len(doc_files)} PDF/DOC file(s) in the table")
    print(f"Extracted metadata: {metadata}")

    # Collect all extracted files
    all_extracted_files = []

    # Download and extract files if requested
    if extract_zips:
        # Process ZIP files first if available
        if zip_files:
            for zip_info in zip_files:
                print(f"Processing ZIP: {zip_info['name']}")
                extraction_result = download_and_extract_zip(zip_info["url"])
                if extraction_result.get("success") and extraction_result.get("files"):
                    # Add all files from this ZIP to the main array
                    all_extracted_files.extend(extraction_result["files"])

        # If no ZIP files, process individual PDF/DOC files
        elif doc_files:
            print("No ZIP files to process. Processing individual PDF/DOC files...")
            for doc_info in doc_files:
                print(f"Processing document: {doc_info['name']}")
                extraction_result = download_and_extract_document(
                    doc_info["url"],
                    doc_info["type"]
                )
                if extraction_result.get("success") and extraction_result.get("file"):
                    # Add the extracted file to the main array
                    all_extracted_files.append(extraction_result["file"])

    return {
        "zip_urls": zip_urls,
        "doc_urls": doc_urls,
        "extracted_files": all_extracted_files,
        "metadata": metadata
    }


def fetch_with_playwright_2captcha_puc(url, wait_time=30, extract_zips=True):
    """
    Synchronous wrapper for playwright_2captcha_fetch_puc.
    Properly handles event loop conflicts when called from Flask.
    Also processes ZIP and document files from the scraped HTML.

    Args:
        url: URL to scrape
        wait_time: Wait time for page load
        extract_zips: Whether to extract files found in the table

    Returns:
        If extract_zips=True: Dictionary with extracted files and metadata
        If extract_zips=False: HTML content as string
    """
    try:
        # First try without nest_asyncio
        html_content = asyncio.run(
            playwright_2captcha_fetch_puc(url, wait_time))
    except RuntimeError as e:
        # If we get a RuntimeError about event loop, apply nest_asyncio
        import nest_asyncio
        nest_asyncio.apply()
        # Now try again with nest_asyncio applied
        html_content = asyncio.run(
            playwright_2captcha_fetch_puc(url, wait_time))

    # If HTML is empty or None, return appropriate structure
    if not html_content:
        if not extract_zips:
            return html_content
        else:
            return {
                "zip_urls": [],
                "doc_urls": [],
                "extracted_files": [],
                "metadata": {}
            }

    # If extract_zips is False, just return HTML
    if not extract_zips:
        return html_content

    # Process files (ZIP or PDF/DOC) and return simplified structure
    return process_puc_documents(html_content, url, extract_zips=True)


if __name__ == "__main__":
    result = fetch_with_playwright_2captcha_puc(TARGET_URL, extract_zips=True)

    if isinstance(result, dict):
        # Result includes file extraction info
        print(f"\n{'='*60}")
        print(f"EXTRACTION RESULTS")
        print(f"{'='*60}")

        # Display metadata
        metadata = result.get("metadata", {})
        if metadata:
            print(f"\nMetadata:")
            for key, value in metadata.items():
                print(f"  {key}: {value}")

        # Display ZIP file URLs
        zip_urls = result.get("zip_urls", [])
        if zip_urls:
            print(f"\nFound {len(zip_urls)} ZIP file(s):")
            for idx, url in enumerate(zip_urls, 1):
                print(f"  {idx}. {url}")

        # Display DOC file URLs
        doc_urls = result.get("doc_urls", [])
        if doc_urls:
            print(f"\nFound {len(doc_urls)} PDF/DOC file(s):")
            for idx, url in enumerate(doc_urls, 1):
                print(f"  {idx}. {url}")

        # Display extracted files with text
        extracted_files = result.get("extracted_files", [])
        if extracted_files:
            print(f"\nExtracted {len(extracted_files)} file(s):")
            for idx, file_info in enumerate(extracted_files, 1):
                file_name = file_info.get("name", "Unknown")
                file_type = file_info.get("type", "unknown")
                text_length = file_info.get("text_length", 0) or len(
                    file_info.get("text", ""))

                print(f"\n  File {idx}: {file_name}")
                print(f"    Type: {file_type}")
                print(f"    Text length: {text_length} characters")

                # Show page count for PDFs
                if file_info.get("page_count"):
                    print(f"    Pages: {file_info.get('page_count')}")

                # Show preview of extracted text
                text = file_info.get("text", "")
                if text:
                    preview = text[:200] + "..." if len(text) > 200 else text
                    print(f"    Text preview: {preview}")

                # Show error if any
                if file_info.get("error"):
                    print(f"    Error: {file_info.get('error')}")

        print(f"\n{'='*60}\n")
    else:
        # Just HTML content
        print("Received HTML content only (no file extraction)")
